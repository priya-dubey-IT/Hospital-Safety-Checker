"""
Service for generating patient medical reports in Word format
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
from typing import Dict

class ReportService:
    """Service for creating and managing patient medical reports"""
    
    @staticmethod
    def create_report_document(report_data: Dict) -> str:
        """
        Create a Word document for a patient report
        
        Args:
            report_data: Dictionary containing report information
            
        Returns:
            Path to the created Word document
        """
        # Create document
        doc = Document()
        
        # Add hospital header
        header = doc.add_heading('HOSPITAL SAFETY CHECKER', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.runs[0]
        header_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Add subtitle
        subtitle = doc.add_paragraph('Patient Medical Report')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.bold = True
        
        doc.add_paragraph()  # Spacing
        
        # Add patient photo if available (Restored only for patient)
        if report_data.get('patient_photo'):
            try:
                p_table = doc.add_table(rows=1, cols=1)
                p_table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell = p_table.rows[0].cells[0]
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                label = para.add_run('Patient Photo\n')
                label.font.bold = True
                
                photo_path = report_data.get('patient_photo')
                if os.path.exists(photo_path):
                    para.add_run().add_picture(photo_path, width=Inches(1.2))
                    doc.add_paragraph()
            except:
                pass

        doc.add_paragraph()  # Spacing after header
        
        doc.add_paragraph()  # Spacing
        
        # Patient Information Section
        doc.add_heading('Patient Information', level=1)
        
        patient_table = doc.add_table(rows=5, cols=2)
        patient_table.style = 'Light Grid Accent 1'
        
        # Format dates properly
        def format_date(date_str):
            if not date_str:
                return 'N/A'
            try:
                # Parse ISO format date
                dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                return dt.strftime('%B %d, %Y at %I:%M %p')
            except:
                return date_str
        
        patient_info = [
            ('Patient Name:', report_data.get('patient_name', 'N/A')),
            ('Age:', report_data.get('patient_age', 'N/A')),
            ('Gender:', report_data.get('patient_gender', 'N/A')),
            ('Registration Date:', format_date(report_data.get('registration_date'))),
            ('Report Completion Date:', format_date(report_data.get('updated_at')))
        ]
        
        for i, (label, value) in enumerate(patient_info):
            patient_table.rows[i].cells[0].text = label
            patient_table.rows[i].cells[1].text = str(value)
            patient_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Spacing
        
        # Medical Information Section
        doc.add_heading('Medical Information', level=1)
        
        # Symptoms
        doc.add_heading('Symptoms:', level=2)
        symptoms_para = doc.add_paragraph(report_data.get('symptoms', 'No symptoms recorded'))
        symptoms_para.style = 'Normal'
        
        # Diagnosis
        doc.add_heading('Diagnosis:', level=2)
        diagnosis_para = doc.add_paragraph(report_data.get('diagnosis', 'No diagnosis recorded'))
        diagnosis_para.style = 'Normal'
        
        # Treatment
        doc.add_heading('Treatment Plan:', level=2)
        treatment_para = doc.add_paragraph(report_data.get('treatment', 'No treatment plan recorded'))
        treatment_para.style = 'Normal'
        
        # Medications
        doc.add_heading('Medications:', level=2)
        medications_para = doc.add_paragraph(report_data.get('medications', 'No medications prescribed'))
        medications_para.style = 'Normal'
        
        # Additional Notes
        if report_data.get('notes'):
            doc.add_heading('Additional Notes:', level=2)
            notes_para = doc.add_paragraph(report_data.get('notes'))
            notes_para.style = 'Normal'
        
        doc.add_paragraph()  # Spacing
        doc.add_paragraph()  # Spacing
        
        # Doctor Signature Section
        doc.add_heading('Attending Physician', level=1)
        
        doctor_table = doc.add_table(rows=2, cols=2)
        doctor_table.style = 'Light Grid Accent 1'
        
        doctor_info = [
            ('Doctor Name:', report_data.get('doctor_name', 'N/A')),
            ('Date:', datetime.now().strftime('%B %d, %Y'))
        ]
        
        for i, (label, value) in enumerate(doctor_info):
            doctor_table.rows[i].cells[0].text = label
            doctor_table.rows[i].cells[1].text = str(value)
            doctor_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph('_' * 80)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = doc.add_paragraph('This is an official medical report from Hospital Safety Checker')
        footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text.runs[0].font.size = Pt(8)
        footer_text.runs[0].font.italic = True
        
        # Save document
        reports_dir = os.path.join(os.getcwd(), "reports")
        os.makedirs(reports_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        patient_name_safe = report_data.get('patient_name', 'patient').replace(' ', '_')
        filename = f"report_{patient_name_safe}_{timestamp}.docx"
        filepath = os.path.join(reports_dir, filename)
        
        doc.save(filepath)
        return filepath
